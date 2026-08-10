"""
Nova File Manager — the filesystem engine behind voice file commands.

Pure functions, zero LLM, zero side effects at import. `file_intents.py` owns
the natural-language layer; this module only knows about paths.

Search is three-pass and TCC-tolerant:
  1. Spotlight literal   `mdfind -name "<query>"`
  2. Spotlight AND-tokens `kMDItemFSName == "*t*"cd && …`
  3. Direct filesystem walk of the user's common folders

Pass 3 matters because Spotlight results are filtered per-process by macOS
privacy: running inside Nova.app without Full Disk Access, `mdfind` can return
nothing for a file that plainly exists on the Desktop. The walk hits the
filesystem directly, so it either works or raises PermissionError — which we
record and surface, rather than reporting a false "I couldn't find it".

Safety, in layers, because these operations touch Nicholas's real files:
  * Nova's own project + data directories are PROTECTED — never surfaced as a
    search candidate, and refused again at the filesystem call itself.
  * Move and copy NEVER overwrite. An existing destination is an error.
  * There is deliberately NO delete. See `file_intents` — Nova declines and
    says why, rather than falling through to an LLM that might claim it
    deleted something.
"""

from __future__ import annotations

import logging
import os
import re
import shutil
import subprocess
from pathlib import Path
from typing import Optional

log = logging.getLogger("nova.files")

_HOME = Path.home()


# ═══════════════════════════════════════════════════════════════════════════
# Protected paths — Nova's own guts are off limits to voice file operations
# ═══════════════════════════════════════════════════════════════════════════
# Resolved at import from this file's location, plus NOVA_DATA_DIR (set by the
# Swift BackendManager) and the Application Support directory that holds the
# memory DB and RAG index. Nova must never move, rename, or read out its own
# config, database, or model weights because someone said "move my file".
_PROJECT_DIR = Path(__file__).resolve().parent


def _protected_roots() -> list[Path]:
    roots = [_PROJECT_DIR, _PROJECT_DIR.parent]  # nova_backend/ and the repo root
    env = os.environ.get("NOVA_DATA_DIR", "").strip()
    if env:
        try:
            roots.append(Path(env).expanduser().resolve())
        except Exception:
            pass
    roots.append(_HOME / "Library" / "Application Support" / "Nova")
    out: list[Path] = []
    for r in roots:
        if r not in out:
            out.append(r)
    return out


def is_protected_path(filepath: str) -> bool:
    """True if the path lives inside Nova's project or data directories."""
    try:
        p = Path(filepath).expanduser().resolve()
    except Exception:
        return False
    for root in _protected_roots():
        try:
            p.relative_to(root)
            return True
        except ValueError:
            continue
    return False


# ═══════════════════════════════════════════════════════════════════════════
# Exclusions — what "my file" never means
# ═══════════════════════════════════════════════════════════════════════════
_EXCLUDED_EXTENSIONS = {
    ".db", ".sqlite", ".sqlite3", ".sqlite-journal", ".db-journal",
    ".db-shm", ".db-wal", ".sqlite-shm", ".sqlite-wal",
    ".pyc", ".pyo", ".pyd", ".so", ".dylib", ".a", ".o", ".bundle",
    ".log", ".lock", ".pid", ".cache", ".tmp", ".temp", ".bak", ".swp",
    ".class", ".jar", ".gguf", ".onnx", ".bin", ".pt", ".ckpt", ".safetensors",
    ".ds_store", ".plist", ".tvdb", ".musicdb", ".itl", ".itc", ".itdb",
}

_EXCLUDED_PATH_PREFIXES = (
    "/System/", "/Library/", "/private/", "/usr/", "/bin/", "/sbin/", "/opt/",
    "/Applications/",
    str(_HOME / "Library") + "/",
)

_EXCLUDED_PATH_FRAGMENTS = (
    "/.Trash/", "/.Spotlight-", "/.DocumentRevisions-", "/.fseventsd/",
    "/node_modules/", "/.git/", "/.venv/", "/__pycache__/",
    "/build/", "/dist/", "/target/",
    ".tvlibrary/", ".musiclibrary/", ".photoslibrary/", ".imovielibrary/",
    ".aplibrary/", ".app/Contents/", ".bundle/Contents/", ".framework/",
)

_TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".rst", ".csv", ".tsv", ".json",
    ".xml", ".yaml", ".yml", ".ini", ".cfg", ".conf", ".toml",
    ".py", ".js", ".ts", ".tsx", ".jsx", ".html", ".css", ".sh", ".swift",
    ".c", ".cpp", ".h", ".hpp", ".java", ".rb", ".go", ".rs",
}
_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".heic", ".webp", ".tiff"}


def _is_excluded(path: str) -> bool:
    if is_protected_path(path):
        return True
    if any(path.startswith(p) for p in _EXCLUDED_PATH_PREFIXES):
        return True
    if any(frag in path for frag in _EXCLUDED_PATH_FRAGMENTS):
        return True
    if Path(path).suffix.lower() in _EXCLUDED_EXTENSIONS:
        return True
    # Anything under a hidden directory in the user's home (~/.cache, ~/.config)
    rel = path[len(str(_HOME)):] if path.startswith(str(_HOME)) else path
    return any(part.startswith(".") for part in rel.split(os.sep) if part)


# ═══════════════════════════════════════════════════════════════════════════
# Query tokenizing
# ═══════════════════════════════════════════════════════════════════════════
# Words that ride along in a spoken request but are not part of the filename.
# Stripping them is what lets "move the budget file from my desktop into
# documents" search for "budget" and nothing else.
_QUERY_STOPWORDS = frozenset({
    # articles / pronouns / glue
    "a", "an", "the", "my", "your", "our", "their", "this", "that", "these",
    "those", "it", "its", "one", "ones", "from", "to", "in", "on", "into",
    "onto", "at", "with", "and", "or", "of", "for", "me", "us", "is", "are",
    "was", "some", "any", "all", "please", "can", "you", "could", "would",
    "should", "do", "does", "just", "up", "over", "there", "here", "where",
    "what", "which", "hey", "nova",
    # generic file nouns — they signal a file request but never identify one
    "file", "files", "document", "documents", "doc", "docs", "folder",
    "directory", "called", "named", "titled", "image", "images", "picture",
    "pictures", "photo", "photos", "screenshot", "screenshots", "spreadsheet",
    "presentation", "slides", "pdf", "pdfs", "thing", "stuff",
    # location words — the user is pointing at a folder, not naming a file
    "desktop", "downloads", "download", "pictures", "music", "movies",
    "videos", "home", "trash", "icloud", "drive",
    # action verbs that leak in when the whole utterance is used as the query
    "move", "moved", "moving", "rename", "renamed", "renaming", "find",
    "finds", "finding", "locate", "located", "locating", "search", "look",
    "show", "summarize", "summarise", "summary", "describe", "read", "open",
    "put", "send", "transfer", "copy", "copied", "duplicate", "get", "grab",
    "pull", "bring", "tell", "about", "inside", "contents", "content",
    # Nova's own vocabulary — never let Nova search for itself
    "nova", "memory", "config", "settings", "database", "log", "logs", "db",
})


def tokenize_query(q: str) -> list[str]:
    """Filler-stripped, lowercase tokens of at least 2 characters.

    Returning an EMPTY list is meaningful: it means the utterance named no
    distinctive file at all ("open my photos"), and the caller should decline
    the intent rather than searching for nothing.
    """
    tokens = [t for t in re.split(r"[^A-Za-z0-9]+", (q or "").lower()) if t]
    return [t for t in tokens if len(t) >= 2 and t not in _QUERY_STOPWORDS]


# ═══════════════════════════════════════════════════════════════════════════
# Search
# ═══════════════════════════════════════════════════════════════════════════
def _run_mdfind(args: list[str]) -> list[str]:
    try:
        result = subprocess.run(args, capture_output=True, text=True, timeout=8)
    except Exception as exc:
        log.warning(f"mdfind failed: {exc}")
        return []
    return [ln.strip() for ln in (result.stdout or "").splitlines() if ln.strip()]


def _filter_and_rank(raw: list[str], limit: int) -> list[str]:
    """Keep real, user-relevant files; most recently modified first."""
    seen: set[str] = set()
    candidates: list[tuple[float, str]] = []
    for path in raw:
        if path in seen or _is_excluded(path):
            continue
        seen.add(path)
        try:
            if not os.path.isfile(path):
                continue
            mtime = os.path.getmtime(path)
        except OSError:
            continue
        candidates.append((mtime, path))
    candidates.sort(key=lambda x: x[0], reverse=True)
    return [p for _, p in candidates[:limit]]


_WALK_ROOTS: tuple[Path, ...] = (
    _HOME / "Desktop", _HOME / "Downloads", _HOME / "Documents",
    _HOME / "Pictures", _HOME / "Music", _HOME / "Movies",
)
_WALK_MAX_DEPTH = 3
_WALK_MAX_HITS = 50

_walk_permission_errors: list[str] = []


def get_last_permission_errors() -> list[str]:
    """Top-level user folders the last search could not read at all.

    Deliberately narrow. An earlier version reported EVERY PermissionError the
    walk hit, which meant `~/Pictures/Photos Library.photoslibrary` — an app
    bundle that is unreadable on every Mac and holds nothing the user would
    ask for — made Nova answer every "I couldn't find it" with a false alarm
    about granting folder access in System Settings.
    """
    roots = {str(r) for r in _WALK_ROOTS}
    return [p for p in _walk_permission_errors if p in roots]


def _walk_on_error(err: OSError) -> None:
    if isinstance(err, PermissionError):
        _walk_permission_errors.append(str(err.filename or "<unknown>"))


def _walk_search(tokens: list[str]) -> list[str]:
    """Filenames in the user's common folders containing ALL tokens."""
    global _walk_permission_errors
    _walk_permission_errors = []
    if not tokens:
        return []

    lowered = [t.lower() for t in tokens]
    hits: list[str] = []

    for root in _WALK_ROOTS:
        if not root.is_dir():
            try:
                root.stat()
            except PermissionError:
                _walk_permission_errors.append(str(root))
            except OSError:
                pass
            continue

        root_str = str(root)
        try:
            for dirpath, dirnames, filenames in os.walk(
                root_str, followlinks=False, onerror=_walk_on_error
            ):
                if dirpath[len(root_str):].count(os.sep) >= _WALK_MAX_DEPTH:
                    dirnames.clear()
                    continue
                # Prune hidden dirs and app-library bundles before descending.
                # Walking into a .photoslibrary is both useless and the source
                # of the permission errors that used to look like a real
                # privacy problem.
                dirnames[:] = [
                    d for d in dirnames
                    if not d.startswith(".")
                    and not any(frag.strip("/") in d for frag in _EXCLUDED_PATH_FRAGMENTS)
                ]
                for fname in filenames:
                    if fname.startswith("."):
                        continue
                    low = fname.lower()
                    if not all(t in low for t in lowered):
                        continue
                    full = os.path.join(dirpath, fname)
                    if _is_excluded(full):
                        continue
                    hits.append(full)
                    if len(hits) >= _WALK_MAX_HITS:
                        return hits
        except PermissionError:
            _walk_permission_errors.append(root_str)
        except OSError:
            continue
    return hits


def search_file(query: str, limit: int = 5) -> list[str]:
    """Find files by name. Most recently modified first, at most `limit`.

    Deliberately does NOT try each token alone — that produced matches like
    /Applications for a query containing "application".
    """
    global _walk_permission_errors
    q = (query or "").strip()
    if not q:
        return []
    # Clear here, not only in _walk_search: a search that succeeds at pass 1 or
    # 2 never walks, and stale errors from an earlier search must not be read
    # back as a permission problem with THIS one.
    _walk_permission_errors = []

    ranked = _filter_and_rank(_run_mdfind(["mdfind", "-name", q]), limit)
    if ranked:
        log.info(f"search {q!r}: {len(ranked)} via spotlight-literal")
        return ranked

    tokens = tokenize_query(q)

    if tokens:
        and_query = " && ".join(f'kMDItemFSName == "*{t}*"cd' for t in tokens)
        ranked = _filter_and_rank(_run_mdfind(["mdfind", and_query]), limit)
        if ranked:
            log.info(f"search {q!r}: {len(ranked)} via spotlight-tokens {tokens}")
            return ranked

    # Pass 3 — walk. Try the whole query squashed to one token AND the token
    # list; whichever finds more wins.
    literal = re.sub(r"\s+", "", q.lower())
    by_literal = _walk_search([literal]) if len(literal) >= 2 else []
    by_tokens = _walk_search(tokens) if tokens else []
    walked = by_tokens if len(by_tokens) >= len(by_literal) else by_literal
    ranked = _filter_and_rank(walked, limit)
    if ranked:
        log.info(f"search {q!r}: {len(ranked)} via walk")
        return ranked

    denied = get_last_permission_errors()
    log.info(f"search {q!r}: no matches"
             + (f" (permission denied on {denied})" if denied else ""))
    return []


# ═══════════════════════════════════════════════════════════════════════════
# File type + content extraction
# ═══════════════════════════════════════════════════════════════════════════
def get_file_type(filepath: str) -> str:
    """One of image / pdf / docx / text / other."""
    ext = Path(filepath).suffix.lower()
    if ext in _IMAGE_EXTENSIONS:
        return "image"
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext in _TEXT_EXTENSIONS:
        return "text"
    return "other"


def human_size(n_bytes: int) -> str:
    size = float(n_bytes)
    for unit in ("bytes", "KB", "MB", "GB", "TB"):
        if size < 1024 or unit == "TB":
            return f"{int(size)} {unit}" if unit == "bytes" else f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} TB"


def spoken_name(filepath: str) -> str:
    """A filename a TTS engine can read without spelling out punctuation.

    "RMV-RealID-Application-Steps.png" reads as a wall of hyphens; this gives
    "RMV RealID Application Steps" and lets the caller name the type separately.
    """
    stem = Path(filepath).stem
    stem = re.sub(r"[-_.]+", " ", stem)
    stem = re.sub(r"\s+", " ", stem).strip()
    return stem or Path(filepath).name


def spoken_kind(filepath: str) -> str:
    """How to say the file's type out loud: 'a PDF', 'a PNG image', 'a text file'."""
    ext = Path(filepath).suffix.lower().lstrip(".")
    kind = get_file_type(filepath)
    if kind == "pdf":
        return "a PDF"
    if kind == "docx":
        return "a Word document"
    if kind == "image":
        return f"a {ext.upper()} image" if ext else "an image"
    if kind == "text":
        return f"a {ext.upper()} file" if ext else "a text file"
    return f"a {ext.upper()} file" if ext else "a file"


def folder_label(filepath: str) -> str:
    """Short spoken name for the folder a file sits in."""
    parent = Path(filepath).expanduser().parent
    if parent == _HOME:
        return "home"
    return parent.name or str(parent)


_MAX_EXTRACT_CHARS = 6000
_PDF_MAX_PAGES = 12


def extract_text(filepath: str) -> tuple[str, Optional[str]]:
    """Pull readable text out of a file for summarization.

    Returns (text, problem). `problem` is a spoken explanation when there is
    nothing to read — a missing library, an image, an empty file — so the
    caller can say why instead of inventing a summary.
    """
    kind = get_file_type(filepath)

    if kind == "image":
        return "", f"{spoken_name(filepath)} is an image, and I can't see inside images yet."

    if kind == "text":
        try:
            with open(filepath, "rb") as fh:
                raw = fh.read(_MAX_EXTRACT_CHARS * 2)
        except Exception as exc:
            log.warning(f"read failed for {filepath}: {exc}")
            return "", "I couldn't read that file."
        text = raw.decode("utf-8", errors="replace").strip()
        if not text:
            return "", f"{spoken_name(filepath)} is empty."
        return text[:_MAX_EXTRACT_CHARS], None

    if kind == "pdf":
        try:
            import pypdf
        except Exception:
            return "", "I can't read PDFs right now, the PDF library isn't installed."
        try:
            reader = pypdf.PdfReader(filepath)
            pages = [(pg.extract_text() or "") for pg in reader.pages[:_PDF_MAX_PAGES]]
        except Exception as exc:
            # Never speak the raw exception: pypdf and python-docx both put the
            # full absolute path in the message, and Nova reads its output aloud.
            log.warning(f"PDF read failed for {filepath}: {exc}")
            return "", "I couldn't open that PDF, it looks damaged."
        text = "\n".join(pages).strip()
        if not text:
            return "", (f"{spoken_name(filepath)} is a PDF with no selectable text, "
                        "it's probably a scan.")
        return text[:_MAX_EXTRACT_CHARS], None

    if kind == "docx":
        try:
            from docx import Document
        except Exception:
            return "", ("I can't read Word documents yet, the python-docx library "
                        "isn't installed.")
        try:
            doc = Document(filepath)
            parts = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" ".join(c.text for c in row.cells))
        except Exception as exc:
            log.warning(f"docx read failed for {filepath}: {exc}")
            return "", "I couldn't open that Word document, it looks damaged."
        text = "\n".join(parts).strip()
        if not text:
            return "", f"{spoken_name(filepath)} is empty."
        return text[:_MAX_EXTRACT_CHARS], None

    return "", f"I don't know how to read {spoken_kind(filepath)}."


def image_dimensions(filepath: str) -> Optional[tuple[int, int]]:
    """Pixel size via macOS `sips` — no extra dependency."""
    try:
        out = subprocess.run(
            ["sips", "-g", "pixelWidth", "-g", "pixelHeight", filepath],
            capture_output=True, text=True, timeout=6,
        ).stdout
    except Exception:
        return None
    w = re.search(r"pixelWidth:\s*(\d+)", out)
    h = re.search(r"pixelHeight:\s*(\d+)", out)
    return (int(w.group(1)), int(h.group(1))) if w and h else None


# ═══════════════════════════════════════════════════════════════════════════
# Operations
# ═══════════════════════════════════════════════════════════════════════════
def _prepare_transfer(source_path: str, destination_path: str, verb: str
                      ) -> tuple[bool, str | tuple[Path, Path]]:
    """Shared validation for move and copy. Returns (False, reason) or
    (True, (src, dst))."""
    src = Path(source_path).expanduser()
    dst = Path(destination_path).expanduser()
    if is_protected_path(str(src)):
        return False, f"that file is part of Nova itself, so I won't {verb} it"
    if is_protected_path(str(dst)):
        return False, "that destination is inside Nova's own files"
    if not src.exists():
        return False, "I couldn't find that file anymore"
    if dst.is_dir():
        dst = dst / src.name
    if dst.resolve() == src.resolve():
        return False, "it's already there"
    if dst.exists():
        return False, f"there's already a file called {dst.name} there"
    return True, (src, dst)


def move_file(source_path: str, destination_path: str) -> tuple[bool, str]:
    """Move a file. Never overwrites, never touches Nova's own directories.
    On success the message is the new absolute path."""
    try:
        ok, payload = _prepare_transfer(source_path, destination_path, "move")
        if not ok:
            return False, str(payload)
        src, dst = payload  # type: ignore[misc]
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(src), str(dst))
        return True, str(dst)
    except Exception as exc:
        log.warning(f"move failed {source_path} -> {destination_path}: {exc}")
        return False, "the move failed"


def copy_file(source_path: str, destination_path: str) -> tuple[bool, str]:
    """Copy a file, leaving the original in place. Same guards as move."""
    try:
        ok, payload = _prepare_transfer(source_path, destination_path, "copy")
        if not ok:
            return False, str(payload)
        src, dst = payload  # type: ignore[misc]
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(str(src), str(dst))
        return True, str(dst)
    except Exception as exc:
        log.warning(f"copy failed {source_path} -> {destination_path}: {exc}")
        return False, "the copy failed"


def rename_file(filepath: str, new_name: str) -> tuple[bool, str]:
    """Rename in place, keeping the directory and (if unspecified) the
    extension. On success the message is the new absolute path."""
    try:
        src = Path(filepath).expanduser()
        if is_protected_path(str(src)):
            return False, "that file is part of Nova itself, so I won't rename it"
        if not src.exists():
            return False, "I couldn't find that file anymore"
        new = (new_name or "").strip().rstrip(".")
        if not new:
            return False, "I didn't catch the new name"
        # A spoken name must never become a path.
        new = new.replace("/", " ").replace("\\", " ").strip()
        if not new:
            return False, "that name isn't usable for a file"
        target = src.with_name(new)
        if not target.suffix and src.suffix:
            target = target.with_suffix(src.suffix)
        if target == src:
            return False, "that's already its name"
        if target.exists():
            return False, f"there's already a file called {target.name} there"
        src.rename(target)
        return True, str(target)
    except Exception as exc:
        log.warning(f"rename failed {filepath} -> {new_name}: {exc}")
        return False, "the rename failed"


def reveal_in_finder(filepath: str) -> bool:
    """Open a Finder window with the file selected."""
    try:
        subprocess.run(["open", "-R", str(Path(filepath).expanduser())],
                       capture_output=True, timeout=8)
        return True
    except Exception as exc:
        log.warning(f"reveal failed: {exc}")
        return False


def open_file(filepath: str) -> bool:
    """Open a file in its default application."""
    try:
        subprocess.run(["open", str(Path(filepath).expanduser())],
                       capture_output=True, timeout=8)
        return True
    except Exception as exc:
        log.warning(f"open failed: {exc}")
        return False


# ═══════════════════════════════════════════════════════════════════════════
# Destination resolution
# ═══════════════════════════════════════════════════════════════════════════
_COMMON_DESTINATIONS = {
    "desktop": _HOME / "Desktop",
    "downloads": _HOME / "Downloads", "download": _HOME / "Downloads",
    "documents": _HOME / "Documents", "document": _HOME / "Documents",
    "docs": _HOME / "Documents",
    "pictures": _HOME / "Pictures", "picture": _HOME / "Pictures",
    "photos": _HOME / "Pictures",
    "music": _HOME / "Music",
    "movies": _HOME / "Movies", "videos": _HOME / "Movies",
    "home": _HOME, "home folder": _HOME,
}


# Where to look for a folder the user names that isn't one of the standards
# ("move it to my projects folder").
_DEST_SEARCH_ROOTS: tuple[Path, ...] = (
    _HOME, _HOME / "Documents", _HOME / "Desktop", _HOME / "Downloads",
)


# ═══════════════════════════════════════════════════════════════════════════
# Folder listing
# ═══════════════════════════════════════════════════════════════════════════
# "What's in my Documents folder?" had NO handler, so it fell all the way to
# the LLM, which cheerfully invented a plausible answer ("a few files") for a
# folder it had never looked at. Reading a directory is trivially knowable —
# there is no excuse for guessing it.
_FOLDER_ALIASES: dict[str, Path] = {
    "desktop": _HOME / "Desktop",
    "documents": _HOME / "Documents",
    "document": _HOME / "Documents",
    "downloads": _HOME / "Downloads",
    "download": _HOME / "Downloads",
    "pictures": _HOME / "Pictures",
    "photos": _HOME / "Pictures",
    "music": _HOME / "Music",
    "movies": _HOME / "Movies",
    "videos": _HOME / "Movies",
    "home": _HOME,
    "home folder": _HOME,
    "applications": Path("/Applications"),
    "apps": Path("/Applications"),
}


def resolve_folder(spoken: Optional[str]) -> Optional[Path]:
    """Map a spoken folder name to a real directory, or None."""
    if not spoken:
        return None
    key = re.sub(r"\s+", " ", str(spoken).lower().strip().strip(".?!,;:"))
    for prefix in ("the ", "my ", "in ", "inside "):
        if key.startswith(prefix):
            key = key[len(prefix):]
    key = re.sub(r"\s+(folder|directory)$", "", key).strip()
    return _FOLDER_ALIASES.get(key)


def list_folder(path: Path, max_names: int = 6) -> dict:
    """Real contents of a directory. Never guesses, never raises.

    Hidden files and Nova's own protected paths are excluded, matching what
    the user would see in Finder. Returns counts plus a few representative
    names so the caller can speak something concrete.
    """
    out: dict = {"path": str(path), "label": folder_label(str(path) + "/x"),
                 "folders": [], "files": [], "n_folders": 0, "n_files": 0,
                 "error": None}
    try:
        entries = sorted(path.iterdir(), key=lambda p: p.name.lower())
    except PermissionError:
        out["error"] = "permission"
        return out
    except FileNotFoundError:
        out["error"] = "missing"
        return out
    except OSError:
        out["error"] = "unreadable"
        return out

    for entry in entries:
        if entry.name.startswith("."):
            continue
        if is_protected_path(str(entry)):
            continue
        try:
            if entry.is_dir():
                out["folders"].append(entry.name)
            elif entry.is_file():
                if entry.suffix.lower() in _EXCLUDED_EXTENSIONS:
                    continue
                out["files"].append(entry.name)
        except OSError:
            continue

    out["n_folders"] = len(out["folders"])
    out["n_files"] = len(out["files"])
    out["folders"] = out["folders"][:max_names]
    out["files"] = out["files"][:max_names]
    return out


def resolve_destination(spoken_location: Optional[str]) -> Optional[str]:
    """Map a spoken folder name to a real directory, or None if there is no
    such place.

    Returning None matters: a spoken name that resolves to a bare relative
    path ("projects") would otherwise be created silently in whatever the
    working directory happens to be. Nova would rather ask.
    """
    if not spoken_location:
        return None
    s = str(spoken_location).strip()
    if not s:
        return None

    # An explicit path is taken at its word.
    if s.startswith(("/", "~")):
        p = Path(s).expanduser()
        return str(p) if p.is_dir() else None

    key = s.lower().strip().strip(".?!,;:")
    for prefix in ("to the ", "to my ", "into the ", "into my ", "the ", "my ", "to ", "into "):
        if key.startswith(prefix):
            key = key[len(prefix):]
            break
    key = re.sub(r"\s+(?:folder|directory)$", "", key).strip()
    if not key:
        return None
    if key in _COMMON_DESTINATIONS:
        return str(_COMMON_DESTINATIONS[key])

    # A folder the user actually has, matched case-insensitively.
    for root in _DEST_SEARCH_ROOTS:
        try:
            for child in root.iterdir():
                if child.is_dir() and child.name.lower() == key and not child.name.startswith("."):
                    return str(child)
        except OSError:
            continue
    return None
